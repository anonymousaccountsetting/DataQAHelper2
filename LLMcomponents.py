import pandas as pd
import os
from docx import Document
import re
import openai
import json
import requests

class SettingForLLM():

    def set_chatGPT(self, background,key, chatmodel="gpt-3.5-turbo", url="https://api.openai.com/v1/chat/completions"):
        openai.api_key = key
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {openai.api_key}"
        }
        #background=set_for_GPT.set_background(Xcol,ycol,modelname,modelinformation)
        background=background
        messages = [{"role": "system", "content": background}, ]
        return (url, background, chatmodel, headers, messages)

    def set_payload(self, message,GPTmodelname="gpt-3.5-turbo", messages=[]):
        messages.append({"role": "user", "content": message}, )
        payload = {
            "model": GPTmodelname,
            "messages": messages,
            "temperature": 1.0,
            "top_p": 1.0,
            "n": 1,
            "stream": False,
            "presence_penalty": 0,
            "frequency_penalty": 0,
        }
        return (payload, messages)

    def send_response_receive_output(self, URL, headers, payload, messages):
        response = requests.post(URL, headers=headers, json=payload, stream=False)
        print(json.loads(response.content))
        output = json.loads(response.content)["choices"][0]['message']['content']
        messages.append({"role": "assistant", "content": output})
        return (output, messages)

    def question_matching(self, question, content, model, url, headers, messages):
        query = (
            "My question is: " + question + "\n"
            "Please refer to the following question bank and choose the Section number "
            "(for example, if you choose Section 5, please return 5.) that matches the meaning of my question. "
            "Please note that as long as the meaning matches, there is no need for word-for-word correspondence. "
            "My entry may have spelling or grammatical mistakes, please ignore those mistakes. "
            "Returns 0 if no section matches. Only answer an integer as you choose, do not reply with any information "
            "other than the integer, do not reply why you chose the section number, do not reply to your thought process. "
            "Following is the question bank: \n" + content
        )

        print("Sending question matching prompt to OpenAI:")
        print(query)

        # 构建 prompt 并加入历史消息
        messages.append({"role": "user", "content": query})
        payload = {
            "model": model,
            "messages": messages,
            "temperature": 0.2,
            "top_p": 1.0,
            "n": 1,
            "stream": False,
            "presence_penalty": 0,
            "frequency_penalty": 0,
        }

        response = requests.post(url, headers=headers, json=payload)
        print("Question Matching Response:")
        print(response.content)

        try:
            result = json.loads(response.content)
            output = result["choices"][0]["message"]["content"].strip()
        except Exception as e:
            output = "0"
            print(f"❌ Failed to parse OpenAI response: {e}")

        return output


    def save_chat_history_to_docx(self,messages):
        doc = Document()

        for message in messages:
            role = message['role']
            content = message['content']

            if role == 'user':
                doc.add_paragraph('User: ' + content)
            elif role == 'assistant':
                doc.add_paragraph('Assistant: ' + content)
            else:
                doc.add_paragraph(content)
        filename="data_report.docx"
        doc.save(filename)


