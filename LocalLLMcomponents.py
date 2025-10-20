from ollama import chat
from ollama import ChatResponse
from docx import Document

class SettingForOllama():

    def extract_after_think(self,text):
        parts = text.split("</think>", 1)
        return parts[1].strip() if len(parts) > 1 else text.strip()



    def set_chat_background(self, messages,background=[]):
        if background==[]:
            background="You are a data analyst assistant. Your task begins by receiving structured information: a dataset description and the analysis results produced by a fitted data science model. After that, you will be given a series of specific questions about the data. For each question, your first job is to match it to one of several predefined abstract questions and return the ID of the most relevant abstract question. Then, based on the answer generated from a predefined answer template for that abstract question, you will revise and polish the answer to produce a concrete, fluent answer that directly addresses the specific question."
        else:
            background="You are a data analyst assistant. Your task begins by receiving structured information: a dataset description and the analysis results produced by a fitted data science model. After that, you will be given a series of specific questions about the data. For each question, your first job is to match it to one of several predefined abstract questions and return the ID of the most relevant abstract question. Then, based on the answer generated from a predefined answer template for that abstract question, you will revise and polish the answer to produce a concrete, fluent answer that directly addresses the specific question. Here are the background knowledge about the dataset: " + background

        backhgroundmessages = [{"role": "system", "content": background}, ]
        print("The background messages is:")
        print(backhgroundmessages)
        messages.append({"role": "system", "content": background})
        return (messages,backhgroundmessages)



    def send_response_receive_output(self, message,backhgroundmessages,messages):
        messages.append({"role": "user", "content": message})
        backhgroundmessages.append({"role": "user", "content": message})
        print("The sended message is:：")
        print(backhgroundmessages)
        response: ChatResponse = chat(model='deepseek-r1:14b', messages=backhgroundmessages)
        backhgroundmessages.remove({"role": "user", "content": message})
        output = SettingForOllama().extract_after_think(response.message.content)
        messages.append({"role": "assistant", "content": output})
        return (output, messages)

    def question_matching(self,question,content):
        query = "My question is: " + question + "\nPlease refer to the following question bank and choose the Section number (for example, if you choose Section 5, please return 5.) that matches the meaning of my question. Please note that as long as the meaning matches, there is no need for word-for-word correspondence. My entry may have spelling or grammatical mistakes, please ignore those mistakes. Returns 0 if no section matches. Only answer an integer as you choose, do not reply with any information other than the integer, do not reply why you chose the section number, do not reply to your thought process. Following is the question bank: \n"+content
        response: ChatResponse = chat(model='deepseek-r1:14b', messages=[{"role": "user", "content": query}])
        print(query)
        print("How the question is matched:：")
        print(response.message.content)
        output = SettingForOllama().extract_after_think(response.message.content)
        return output


    def save_chat_history_to_docx(self,messages):
        print("The messages used for summary is:")
        print(messages)
        doc = Document()

        for message in messages:
            role = message['role']
            content = message['content']

            if role == 'user':
                doc.add_paragraph('User: ' + content)
            elif role == 'assistant':
                doc.add_paragraph('Assistant: ' + content)
            else:
                doc.add_paragraph(content)
        filename="data_report.docx"
        doc.save(filename)


