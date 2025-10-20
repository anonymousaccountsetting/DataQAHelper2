import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from sklearn.impute import SimpleImputer
import os
from docx import Document
import statsmodels.api as sm
from sklearn.discriminant_analysis import LinearDiscriminantAnalysis
from sklearn.linear_model import LogisticRegression, RidgeClassifier,LinearRegression, Ridge, Lasso, BayesianRidge
from sklearn.ensemble import GradientBoostingRegressor, RandomForestRegressor,RandomForestClassifier
from sklearn.tree import DecisionTreeClassifier
from sklearn.svm import SVC
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score,roc_auc_score, roc_curve,r2_score
import numpy as np
import re
import openai
import json
import requests
from jinja2 import Environment, FileSystemLoader
from pycaret import classification,regression
# from langchain.document_loaders import TextLoader
# from langchain.indexes import VectorstoreIndexCreator

# Loading the folder that contains the txt templates

file_loader = FileSystemLoader('./apptemplates')

# Creating a Jinja Environment

env = Environment(loader=file_loader)

# # Loading the Jinja templates from the folder
# For the regression
linearSummary = env.get_template('linearSummary.txt')
linearSummary2 = env.get_template('linearSummary2.txt')
linearSummary3 = env.get_template('linearSummary3.txt')
linearSummary4 = env.get_template('linearSummary4.txt')
linearQuestion = env.get_template('linearQuestionset.txt')

# For Machine Learning Model
MachineLearningLinearModelSummary1=env.get_template('MachineLearningLinearModelSummary1.txt')
MachineLearningLinearModelSummary2=env.get_template('MachineLearningLinearModelSummary2.txt')
MLlinearQuestionSet=env.get_template('MLlinearQuestionSet.txt')
MachineLearningclassifierSummary1=env.get_template('MachineLearningClassifierModelSummary1.txt')
MachineLearningclassifierSummary2=env.get_template('MachineLearningClassifierModelSummary2.txt')
MLclassifierQuestionSet=env.get_template('MLclassifierQuestionSet.txt')

# For the classifier
classifiersummary1=env.get_template('classifiersummary1.txt')
classifiersummary2=env.get_template('classifiersummary2.txt')
classifiersummary3=env.get_template('classifiersummary3.txt')
classifierquestion=env.get_template('classifierQuestionset.txt')



# For ChatGPT
databackground = env.get_template('databackground.txt')
# questionrequest=env.get_template('question_request.txt')
# questionselect=env.get_template('question_select_request.txt')
answerup=env.get_template('answer_upgrade.txt')

# for pycaret
automodelcompare1 = env.get_template('AMC1.txt')
automodelcompare2 = env.get_template('AMC2.txt')

def extract_first_integer(string):
    match = re.search(r'\d+', string)
    if match:
        return int(match.group())
    else:
        return 0

class MyApp:
    def __init__(self):
        self.window = tk.Tk()
        self.window.geometry('800x600')
        self.window.title('My Application')
        self.csv_data = None
        self.independent_vars = None
        self.dependent_var = None
        self.choice = None
        self.criterion = None

    def start(self):
        self.step1()
        self.window.mainloop()

    def step1(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 1: Select a CSV file").pack()
        tk.Button(self.window, text="Select", command=self.load_csv).pack()
        tk.Button(self.window, text="Next", command=self.step2, anchor='e').pack(side='bottom')

    def load_csv(self):
        filename = filedialog.askopenfilename(filetypes=(("CSV files", "*.csv"),))
        if filename:
            self.csv_data = pd.read_csv(filename)

    def step2(self):
        if self.csv_data is None:
            messagebox.showerror("No File", "Please select a CSV file first.")
            return
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 2: Select independent and dependent variables").pack()

        frame = tk.Frame(self.window)
        frame.pack()

        tk.Label(frame, text="Please select the dependent variable first:").pack(side='left')
        self.dependent_var = ttk.Combobox(frame)
        self.dependent_var.pack(side='left')

        tk.Label(frame, text="Please select the independent variables:").pack(side='right')
        self.independent_vars = tk.Listbox(frame, selectmode='multiple')
        self.independent_vars.pack(side='right')

        headers = self.csv_data.columns.tolist()
        self.dependent_var['values'] = headers
        for header in headers:
            self.independent_vars.insert(tk.END, header)

        tk.Button(self.window, text="Back", command=self.step1, anchor='w').pack(side='bottom')
        tk.Button(self.window, text="Next", command=self.save_selections_and_go_to_step3, anchor='e').pack(
            side='bottom')

    def save_selections_and_go_to_step3(self):
        self.selected_independent_vars = self.get_selected_independent_vars()
        self.selected_dependent_var = self.dependent_var.get()
        self.step3()

    def save_criterion_and_go_to_step4(self):
        if self.criterion.get()=="":
            if self.choice.get()==3:
                self.selected_criterion="r2"
            elif self.choice.get()==4:
                self.selected_criterion="Accuracy"
            else:
                self.selected_criterion="unnecessary"
        else:
            self.selected_criterion = self.criterion.get()
        if self.choice.get() == 1:
            self.step4_1()
        elif self.choice.get() == 2:
            self.step4_2()
        elif self.choice.get() == 3:
            self.step4_3()
        elif self.choice.get() == 4:
            self.step4_4()

    def step3(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 3: Select a model type and criterion").pack()

        self.choice = tk.IntVar()

        tk.Radiobutton(self.window, text=f"Option 1: Use a regression model to explore the data.", variable=self.choice, value=1).pack()
        tk.Radiobutton(self.window, text=f"Option 2: Use a classifier model to explore the data.", variable=self.choice, value=2).pack()
        tk.Radiobutton(self.window, text=f"Option 3: Automatically fit a best regression model to explore the data.", variable=self.choice, value=3).pack()
        tk.Radiobutton(self.window, text=f"Option 4: Automatically fit a best classifier model to explore the data.", variable=self.choice, value=4).pack()

        tk.Label(self.window, text="Criterion:").pack()
        self.criterion = tk.Entry(self.window, text="Enter criterion for best model:")
        self.criterion.config(text=f"User input: {self.criterion.get()}")
        self.criterion.pack(pady=10)

        tk.Button(self.window, text="Back", command=self.step2, anchor='w').pack(side='bottom')

        tk.Button(self.window, text="Next", command=self.save_criterion_and_go_to_step4, anchor='e').pack(side='bottom')


    def step4_1(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 4: Choose a regression model.").grid()

        self.model_var = tk.IntVar()

        # Use Statsmodels Linear Regression
        self.statsmodels_button = tk.Radiobutton(self.window, text="Use Statsmodels Linear Regression",
                                                 variable=self.model_var, value=0)
        self.statsmodels_button.grid(row=2, column=0, sticky=tk.E+tk.W, padx=10)

        # Use Scikit-learn Linear Regression
        self.sklearn_button = tk.Radiobutton(self.window, text="Use Scikit-learn Linear Regression",
                                             variable=self.model_var, value=1)
        self.sklearn_button.grid(row=3, column=0, sticky=tk.E+tk.W, padx=10)

        # Add Ridge Regression option
        self.ridge_checkbutton = tk.Radiobutton(self.window, text="Use Ridge Regression", variable=self.model_var,
                                                value=2)
        self.ridge_checkbutton.grid(row=4, column=0, sticky=tk.E+tk.W, padx=10)

        self.ridge_alpha_label = tk.Label(self.window, text="Enter Alpha value (default 1.0):")
        self.ridge_alpha_label.grid(row=5, column=0, sticky=tk.E+tk.W, padx=10)
        self.ridge_alpha_entry = tk.Entry(self.window)
        self.ridge_alpha_entry.grid(row=6, column=0, sticky=tk.E+tk.W, padx=10, pady=5)


        # Use GB Regression
        self.GB_button = tk.Radiobutton(self.window, text="Use Gradient Boosting Regression",
                                             variable=self.model_var, value=5)
        self.GB_button.grid(row=7, column=0, sticky=tk.E+tk.W, padx=10)


        # Use RF Regression
        self.RF_button = tk.Radiobutton(self.window, text="Use Random Forest Regression",
                                             variable=self.model_var, value=6)
        self.RF_button.grid(row=8, column=0, sticky=tk.E+tk.W, padx=10)



        # Add Lasso Regression option
        self.lasso_checkbutton = tk.Radiobutton(self.window, text="Use Lasso Regression", variable=self.model_var,
                                                value=3)
        self.lasso_checkbutton.grid(row=2, column=1, sticky=tk.E+tk.W, padx=10)

        self.lasso_alpha_label = tk.Label(self.window, text="Enter Alpha value (default 1.0):")
        self.lasso_alpha_label.grid(row=3, column=1, sticky=tk.E+tk.W, padx=10)
        self.lasso_alpha_entry = tk.Entry(self.window)
        self.lasso_alpha_entry.grid(row=4, column=1, sticky=tk.E+tk.W, padx=10, pady=5)

        # Add Bayesian Ridge Regression option
        self.bayesian_ridge_checkbutton = tk.Radiobutton(self.window, text="Use Bayesian Ridge Regression",
                                                         variable=self.model_var, value=4)
        self.bayesian_ridge_checkbutton.grid(row=5, column=1, sticky=tk.E+tk.W, padx=10)

        self.bayesian_ridge_label = tk.Label(self.window, text="Enter two Alpha values, and Lambda values:")
        self.bayesian_ridge_label.grid(row=6, column=1, sticky=tk.E+tk.W, padx=10)
        self.alpha_1_entry = tk.Entry(self.window, width=8)
        self.alpha_1_entry.grid(row=7, column=1, sticky=tk.E+tk.W, padx=10)
        self.alpha_2_entry = tk.Entry(self.window, width=8)
        self.alpha_2_entry.grid(row=7, column=1, padx=(100, 0), sticky=tk.E+tk.W)
        self.lambda_1_entry = tk.Entry(self.window, width=8)
        self.lambda_1_entry.grid(row=7, column=1, padx=(190, 0), sticky=tk.E+tk.W)
        self.lambda_2_entry = tk.Entry(self.window, width=8)
        self.lambda_2_entry.grid(row=7, column=1, padx=(280, 0), sticky=tk.E+tk.W)

        tk.Button(self.window, text="Next", command=self.perform_regression, anchor='e').grid()


    def step4_2(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 4: Choose a classifier model.").grid()

        self.model_var = tk.IntVar()

        # Use Logistic Regression
        self.Logistic_button = tk.Radiobutton(self.window, text="Use Logistic Regression",
                                                 variable=self.model_var, value=0)
        self.Logistic_button.grid(row=2, column=0, sticky=tk.E+tk.W, padx=10)

        # Use Linear Discriminant Analysis
        self.LDA_button = tk.Radiobutton(self.window, text="Use Linear Discriminant Analysis",
                                             variable=self.model_var, value=1)
        self.LDA_button.grid(row=3, column=0, sticky=tk.E+tk.W, padx=10)

        # Add SVM - Linear Kernel option
        self.SVM_checkbutton = tk.Radiobutton(self.window, text="Use SVM - Linear Kernel", variable=self.model_var,
                                                value=2)
        self.SVM_checkbutton.grid(row=4, column=0, sticky=tk.E+tk.W, padx=10)

        # Add Ridge Classifier option
        self.RC_checkbutton = tk.Radiobutton(self.window, text="Use Ridge Classifier", variable=self.model_var,
                                                value=3)
        self.RC_checkbutton.grid(row=5, column=0, sticky=tk.E+tk.W, padx=10)

        # Add Random Forest Classifier option
        self.RF_checkbutton = tk.Radiobutton(self.window, text="Use Random Forest Classifier", variable=self.model_var,
                                                value=4)
        self.RF_checkbutton.grid(row=6, column=0, sticky=tk.E+tk.W, padx=10)

        # Add Ridge Classifier option
        self.DT_checkbutton = tk.Radiobutton(self.window, text="Use Decision Tree Classifier", variable=self.model_var,
                                                value=5)
        self.DT_checkbutton.grid(row=7, column=0, sticky=tk.E+tk.W, padx=10)

        tk.Button(self.window, text="Next", command=self.perform_classifier, anchor='e').grid()


    def step4_3(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 4: Here is the regression model selected for you.").pack()

        X = self.csv_data[self.selected_independent_vars]
        y = self.csv_data[self.selected_dependent_var]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        dataset = pd.concat([X_train, y_train], axis=1)

        reg = regression.setup(data=dataset, target=self.selected_dependent_var)
        exclude=['catboost','lightgbm','et','ada','xgboost','llar','lar','huber','dt','omp','par','en','knn','dummy']
        best_model = regression.compare_models(exclude=exclude, n_select=1, sort=self.selected_criterion)
        comapre_results = regression.pull()
        self.p_values = sm.OLS(y, sm.add_constant(X)).fit().pvalues
        self.coefficients = np.append(best_model.intercept_, best_model.coef_)
        self.r_squared = r2_score(y_test, best_model.predict(X_test))
        data_dict = {'Coefficients': self.coefficients[1:]}
        data_dict['P-values'] = self.p_values[1:]
        self.coef_pval_df = pd.DataFrame(data_dict, index=self.selected_independent_vars)
        self.coef_pval_df.index.name = "Xcol"
        self.coef_pval_df = self.coef_pval_df.reset_index()
        self.modeldetail=str(best_model)
        self.more_readable_model_name()

        modelcomparestory = automodelcompare1.render(best=self.modelname, detail=self.modeldetail, n_select=1, sort=self.selected_criterion)

        print(self.r_squared)
        print(comapre_results)
        print(modelcomparestory)
        self.modelchoose_label = tk.Text(self.window, width=100, height=40)
        self.modelchoose_label.delete(1.0, tk.END)
        self.modelchoose_label.insert(tk.END, modelcomparestory)
        self.modelchoose_label.pack(pady=10)
        tk.Button(self.window, text="Next", command=self.step5, anchor='e').pack()

    def step4_4(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 4: Here is the classifier model selected for you.").pack()

        X = self.csv_data[self.selected_independent_vars]
        y = self.csv_data[self.selected_dependent_var]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        dataset = pd.concat([X_train, y_train], axis=1)
        clf = classification.setup(data=dataset, target=self.selected_dependent_var)
        exclude = ['qda','knn','lightgbm','et','catboost','xgboost','gbc','ada','dummy']

        best_model = classification.compare_models(exclude=exclude, n_select=1, sort=self.selected_criterion)
        comapre_results = classification.pull()
        print(comapre_results)

        self.coefficients=best_model.coef_
        target_names = best_model.classes_

        if len(target_names) == 2 and best_model.coef_.shape[0] == 1:
            # Special handling for binary classification case
            self.coeff_df = pd.DataFrame(best_model.coef_, columns=self.selected_independent_vars, index=[target_names[0]])
        else:
            # The usual case for multiclass classification
            self.coeff_df = pd.DataFrame(best_model.coef_, columns=self.selected_independent_vars, index=target_names)

        # Make predictions on the test set

        print(self.coeff_df)
        y_pred = best_model.predict(X_test)
        # Calculate and output the accuracy
        self.accuracy = accuracy_score(y_test, y_pred)
        self.modeldetail = str(best_model)
        self.more_readable_model_name()
        modelcomparestory = automodelcompare1.render(best=self.modelname, detail=self.modeldetail, n_select=1, sort=self.selected_criterion)

        print(modelcomparestory)
        self.modelchoose_label = tk.Text(self.window, width=100, height=40)
        self.modelchoose_label.delete(1.0, tk.END)
        self.modelchoose_label.insert(tk.END, modelcomparestory)
        self.modelchoose_label.pack(pady=10)
        tk.Button(self.window, text="Next", command=self.step5, anchor='e').pack()

    def perform_regression(self):
        threshold=0.8
        self.csv_data=self.cleanData( self.csv_data, threshold, Xcol=self.selected_independent_vars, ycol=self.selected_dependent_var)
        X = self.csv_data[self.selected_independent_vars]
        y = self.csv_data[self.selected_dependent_var]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

        if self.model_var.get() == 0:
            X = sm.add_constant(X)
            model = sm.OLS(y, X).fit()
            self.coefficients = model.params  # Save slope
            self.p_values = model.pvalues  # Save p-values
            self.r_squared = model.rsquared
        elif self.model_var.get()==1:
            model = LinearRegression().fit(X, y)
            self.coefficients = np.append(model.intercept_, model.coef_)
            self.p_values = sm.OLS(y, sm.add_constant(X)).fit().pvalues
            self.r_squared = model.score(X, y)
        elif self.model_var.get()==2:
            self.p_values = sm.OLS(y, sm.add_constant(X)).fit().pvalues

            alpha = float(self.ridge_alpha_entry.get() or 1.0)
            model = Ridge(alpha=alpha).fit(X_train, y_train)
            self.coefficients = np.append(model.intercept_, model.coef_)
            self.r_squared = r2_score(y_test, model.predict(X_test))
        elif self.model_var.get()==3:
            self.p_values = sm.OLS(y, sm.add_constant(X)).fit().pvalues

            alpha = float(self.lasso_alpha_entry.get() or 1.0)
            model = Lasso(alpha=alpha).fit(X_train, y_train)
            self.coefficients = np.append(model.intercept_, model.coef_)
            self.r_squared = r2_score(y_test, model.predict(X_test))
        elif self.model_var.get()==4:
            self.p_values = sm.OLS(y, sm.add_constant(X)).fit().pvalues

            alpha_1 = float(self.alpha_1_entry.get() or 1e-06)
            alpha_2 = float(self.alpha_2_entry.get() or 1e-06)
            lambda_1 = float(self.lambda_1_entry.get() or 1e-06)
            lambda_2 = float(self.lambda_2_entry.get() or 1e-06)
            model = BayesianRidge(alpha_1=alpha_1, alpha_2=alpha_2, lambda_1=lambda_1, lambda_2=lambda_2).fit(X_train, y_train)
            # Get the coefficients for each predictor
            self.coefficients = np.append(model.intercept_, model.coef_)
            # Calculate R-squared
            y_pred = model.predict(X_test)
            self.r_squared = r2_score(y_test, y_pred)

        elif self.model_var.get()==5:
            self.p_values = sm.OLS(y, sm.add_constant(X)).fit().pvalues

            # Create Gradient Boosting Regressor
            model = GradientBoostingRegressor(random_state=42)
            model.fit(X_train, y_train)

            # Get the coefficients for each predictor
            self.coefficients = np.append(0, model.feature_importances_)
            # Calculate R-squared
            y_pred = model.predict(X_test)
            self.r_squared = r2_score(y_test, y_pred)
            self.r2_train = model.score(X_train, y_train)
            self.r2_test = model.score(X_test, y_test)

        elif self.model_var.get()==6:
            self.p_values = sm.OLS(y, sm.add_constant(X)).fit().pvalues

            # Create Random Forest Regressor
            model = RandomForestRegressor(random_state=42)
            model.fit(X_train, y_train)

            # Get the coefficients for each predictor
            self.coefficients = np.append(0, model.feature_importances_)

            # Calculate R-squared
            y_pred = model.predict(X_test)
            self.r_squared = r2_score(y_test, y_pred)
            self.r2_train = model.score(X_train, y_train)
            self.r2_test = model.score(X_test, y_test)

        # if self.model_var.get() != 0:
        #     print(model.feature_importances_)


        # Create a dataframe with coefficients and p-values (if available) for independent variables only
        data_dict = {'Coefficients': self.coefficients[1:]}
        data_dict['P-values'] = self.p_values[1:]
        self.coef_pval_df = pd.DataFrame(data_dict, index=self.selected_independent_vars)
        self.coef_pval_df.index.name = "Xcol"
        self.coef_pval_df = self.coef_pval_df.reset_index()

        print(data_dict)

        self.step5()

    def perform_classifier(self):
        X = self.csv_data[self.selected_independent_vars]
        y = self.csv_data[self.selected_dependent_var]

        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

        if self.model_var.get() == 0:
            model = LogisticRegression(max_iter=1000)

        elif self.model_var.get()==1:
            model = LinearDiscriminantAnalysis()

        elif self.model_var.get()==2:
            model = SVC(kernel='linear')

        elif self.model_var.get()==3:
            model = RidgeClassifier()

        elif self.model_var.get()==4:
            model = RandomForestClassifier()
        elif self.model_var.get()==5:
            model = DecisionTreeClassifier()

        model.fit(X_train, y_train)

        if self.model_var.get()<4:
            # print(model.feature_importances_)
            self.coefficients=model.coef_
            self.target_class = model.classes_
            target_names = model.classes_

            if len(target_names) == 2 and model.coef_.shape[0] == 1:
                # Special handling for binary classification case
                self.coeff_df = pd.DataFrame(model.coef_, columns=self.selected_independent_vars, index=[target_names[0]])
            else:
                # The usual case for multiclass classification
                self.coeff_df = pd.DataFrame(model.coef_, columns=self.selected_independent_vars, index=target_names)
            # Add pseudo column for index values
            self.coeff_df.insert(0, self.selected_dependent_var, self.coeff_df.index)

            # Reset index
            self.coeff_df.reset_index(drop=True, inplace=True)
            # print(model.coef_)
            # print(self.selected_independent_vars)
            # print(target_names)
        else:
            self.coefficients=model.feature_importances_
            self.coeff_df=pd.DataFrame([model.feature_importances_], columns=self.selected_independent_vars)
        print(self.coefficients)
        print(self.coeff_df)
        # feature_importance = np.mean(np.abs(self.coefficients), axis=0)
        # most_important_x = self.coeff_df.columns[np.argmax(np.abs(feature_importance))]
        # print(np.abs(feature_importance))
        # print(most_important_x)
        y_pred = model.predict(X_test)
        # Calculate and output the accuracy
        self.accuracy = accuracy_score(y_test, y_pred)

        self.train_accuracy = model.score(X_train, y_train)
        self.test_accuracy = model.score(X_test, y_test)
        self.step5()



    def step5(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 5: Question Asking").pack()

        self.user_text_label = tk.Label(self.window, text="Enter a question about the dataset or the model:")
        self.user_text_label.pack(pady=10)

        self.user_text = tk.StringVar()
        self.user_text_entry = tk.Entry(self.window, textvariable=self.user_text, width=50)
        self.user_text_entry.pack(pady=10)

        self.api_key_label = tk.Label(self.window, text="Enter your API key:")
        self.api_key_label.pack(pady=10)

        self.api_key = tk.StringVar()
        self.api_key_entry = tk.Entry(self.window, textvariable=self.api_key, show='*', width=50)
        self.api_key_entry.pack(pady=10)

        self.gpt_model_label = tk.Label(self.window, text="Enter GPT model name (default: gpt-3.5-turbo):")
        self.gpt_model_label.pack(pady=10)

        self.gpt_model_name = tk.StringVar()
        self.gpt_model_entry = tk.Entry(self.window, textvariable=self.gpt_model_name)
        self.gpt_model_entry.pack(pady=10)

        self.background_text_label = tk.Label(self.window, text="Enter background about the dataset:")
        self.background_text_label.pack(pady=10)

        self.background_text = tk.StringVar()
        self.background_text_entry = tk.Entry(self.window, textvariable=self.background_text, width=50)
        self.background_text_entry.pack(pady=10)

        self.target_text_label = tk.Label(self.window, text="Enter target text:")
        self.target_text_label.pack(pady=10)

        self.target_text = tk.StringVar()
        self.target_text_entry = tk.Entry(self.window, textvariable=self.target_text, width=50)
        self.target_text_entry.pack(pady=10)

        tk.Button(self.window, text="Next", command=self.set_background, anchor='e').pack()

    def step6(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 6: Question Answering").pack()

        self.answer_label = tk.Text(self.window, width=100, height=40)
        self.answer_label.delete(1.0, tk.END)
        self.answer_label.insert(tk.END, self.GPTanswer)
        self.answer_label.pack(pady=10)

        self.ask_another_button = tk.Button(self.window, text="Ask another question", command=self.step8)
        self.ask_another_button.pack( pady=10)

    def step8(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 7: Asking another question").pack()

        self.ask_another_question_label = tk.Label(self.window, text="Enter another question for Chat-GPT:")
        self.ask_another_question_label.pack(pady=10)

        self.another_question_text = tk.StringVar()
        self.another_question_entry = tk.Entry(self.window, textvariable=self.another_question_text, width=50)
        self.another_question_entry.pack(pady=10)

        self.next_button = tk.Button(self.window, text="Next", command=self.go_to_step9)
        self.next_button.pack(side=tk.LEFT, pady=10)


    def step9(self):
        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 8: Showing the asnwer").pack()

        self.new_answer_label = tk.Text(self.window, width=100, height=40)
        self.new_answer_label.delete(1.0, tk.END)
        self.new_answer_label.insert(tk.END, self.new_answer)
        self.new_answer_label.pack(pady=10)

        self.ask_another_button = tk.Button(self.window, text="Ask another question", command=self.step8)
        self.ask_another_button.pack(side=tk.LEFT, pady=10)

        self.make_summary=tk.Button(self.window, text="Summary and End", command=self.step10)
        self.make_summary.pack(side=tk.LEFT, pady=10)

    def step10(self):

        for widget in self.window.winfo_children():
            widget.destroy()
        tk.Label(self.window, text="Step 9: Summary").pack()

        question = "Now, to summarize the above information. Write a short text report using the answers to user questions about the dataset and model. Note that this report only focuses on questions raised by users about the dataset and model. And no need to give an explanation of how this conclusion was reached."

        user_target_text = self.target_text.get()
        if user_target_text:
            question=question+" When you summarize, please keep your writing style and format consistent with what follows. \n"+user_target_text
        payload, messages = self.set_payload(question, self.gpt_model_name, self.messages)
        output, messages = self.send_response_receive_output(self.url, self.headers, payload, messages)
        self.messages = messages
        self.summary=output

        self.summary_label = tk.Text(self.window, width=100, height=40)
        self.summary_label.delete(1.0, tk.END)
        self.summary_label.insert(tk.END, self.summary)
        self.summary_label.pack(pady=10)
        self.save_chat_history_to_docx()
        self.again_button = tk.Button(self.window, text="Do it again", command=self.step1)
        self.again_button.pack(side=tk.LEFT, pady=10)

    def go_to_step9(self):

        self.user_question=self.another_question_text.get()
        self.question_number_select()
        if self.choice.get()==1 or self.choice.get()==3:
            self.regression_answer()
            print(self.messages)

        elif self.choice.get()==2 or self.choice.get()==4:
            self.classifier_answer()
            print(self.messages)

        self.new_answer = self.output
        print(self.messages)

        self.step9()

    def set_background(self):
        # user_api_key = self.api_key.get()
        # if not user_api_key:
        #     self.answerbyGPT.config(text="Invalid API key. Please enter a valid API key.")
        #     return
        # key=user_api_key
        key='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'
        os.environ['OPENAI_API_KEY'] = key
        gpt_model_name_input = self.gpt_model_name.get()
        gpt_model_name = gpt_model_name_input if gpt_model_name_input else "gpt-3.5-turbo"
        self.gpt_model_name=gpt_model_name

        if self.choice.get() == 1:
            if self.model_var.get() == 0:
                self.modelname="Statsmodels Linear Regression"
            elif self.model_var.get() == 1:
                self.modelname="Scikit-learn Linear Regression"
            elif self.model_var.get() == 2:
                self.modelname="Ridge Regression"
            elif self.model_var.get() == 3:
                self.modelname="Lasso Regression"
            elif self.model_var.get() == 4:
                self.modelname="Bayesian Ridge Regression"
            elif self.model_var.get() == 5:
                self.modelname = "Gradient Boosting Regressor"
            elif self.model_var.get() == 6:
                self.modelname = "Random Forest Regressor"
            modelinformation=f"The R-squared of the model is: {self.r_squared:.4f}\n"+str(self.coef_pval_df)

        elif self.choice.get() == 2:
            if self.model_var.get() == 0:
                self.modelname = "Logistic Regression"
            elif self.model_var.get() == 1:
                self.modelname = "Linear Discriminant Analysis"
            elif self.model_var.get() == 2:
                self.modelname = "SVM-Linear Kernel"
            elif self.model_var.get() == 3:
                self.modelname = "Ridge Classifier"
            elif self.model_var.get() == 4:
                self.modelname = "Random Forest Classifier"
            elif self.model_var.get() == 5:
                self.modelname = "Decision Tree Classifier"
            modelinformation = f"The accuracy of the model is: {self.accuracy:.4f}\n" + str(self.coeff_df)

        elif self.choice.get() == 3:
            modelinformation = f"The R-squared of the model is: {self.r_squared:.4f}\n" + str(self.coef_pval_df)

        elif self.choice.get() == 4:
            modelinformation = f"The accuracy of the model is: {self.accuracy:.4f}\n" + str(self.coeff_df)

        if self.background_text !="":
            modelinformation=modelinformation+'\n'+self.background_text.get()

        print(modelinformation)

        self.url, self.background, self.chatmodel, self.headers, self.messages=self.set_chatGPT(self.selected_independent_vars, self.selected_dependent_var, self.modelname, modelinformation,key)
        if self.choice.get() == 1 or self.choice.get() ==3:
            self.go_to_step6_1()
        elif self.choice.get() == 2 or self.choice.get() ==4:
            self.go_to_step6_2()


    def go_to_step6_1(self):
        self.user_question=self.user_text.get()

        self.question_number_select()

        self.regression_answer()
        print(self.messages)
        self.GPTanswer=self.output
        self.step6()

    def regression_answer(self):
        section_num = self.section_num
        questions=[]
        answers=[]
        if section_num==1:
            questions = linearQuestion.render(section=1, indeNum=len(self.coef_pval_df['Xcol']),
                                             xcol=self.coef_pval_df['Xcol'], ycol=self.selected_dependent_var)

            answers=linearSummary2.render(modelname=self.modelname,r2=self.r_squared)
        elif section_num==2:
            for index, row in self.coef_pval_df.iterrows():
                question = linearQuestion.render(section=2, xcol=row['Xcol'], ycol=self.selected_dependent_var)
                answer = linearSummary.render(coeff=row['Coefficients'], p=row['P-values'], xcol=row['Xcol'],
                                              ycol=self.selected_dependent_var)
                questions.append(question)
                answers.append(answer)
        elif section_num==3:

            # Update the most_important_x
            questions = linearQuestion.render(section=3, ycol=self.selected_dependent_var)
            most_important_x = self.selected_independent_vars[np.argmax(np.abs(self.coefficients[1:]))]
            # Display X variables with positive and negative slopes
            positive_slopes = [x for i, x in enumerate(self.selected_independent_vars) if self.coefficients[i + 1] > 0]
            negative_slopes = [x for i, x in enumerate(self.selected_independent_vars) if self.coefficients[i + 1] < 0]
            # Display X variables with P-values greater than and less than 0.05
            high_p_values = [x for i, x in enumerate(self.selected_independent_vars) if self.p_values[i + 1] > 0.05]
            low_p_values = [x for i, x in enumerate(self.selected_independent_vars) if self.p_values[i + 1] <= 0.05]
            answers = linearSummary3.render(ss=low_p_values, lenss=len(low_p_values), nss=high_p_values,
                                           lennss=len(high_p_values), pf=positive_slopes, lenpf=len(positive_slopes),
                                           nf=negative_slopes, lennf=len(negative_slopes), ycol=self.selected_dependent_var,
                                           imp=most_important_x)

        elif section_num==4:
            for index, row in self.coef_pval_df.iterrows():
                question = linearQuestion.render(section=4, xcol=row['Xcol'], ycol=self.selected_dependent_var)
                answer = linearSummary4.render(coeff=row['Coefficients'], p=row['P-values'], xcol=row['Xcol'],
                                              ycol=self.selected_dependent_var)
                questions.append(question)
                answers.append(answer)

        elif section_num==5:
            questions = MLlinearQuestionSet.render(section=5, ycol=self.selected_dependent_var)
            most_important_x = self.selected_independent_vars[np.argmax(np.abs(self.coefficients[1:]))]
            answers = MachineLearningLinearModelSummary1.render(ycol=self.selected_dependent_var,Xcol=most_important_x,coeff=np.abs(self.coefficients[1:]))
        elif section_num == 6:
            questions = MLlinearQuestionSet.render(section=6, ycol=self.selected_dependent_var)
            answers = MachineLearningLinearModelSummary2.render(trainR2=self.r2_train, testR2=self.r2_test)
        elif section_num==0:
            questions="There are no matching default questions in the template."
            answers="Please answer the question based on the analysis results."

        print(questions)
        print(answers)
        # default_answer=answerup.render(userquestion=self.user_text.get(),questions=questions,answers=answers)

        default_answer=answerup.render(userquestion=self.user_question,questions=questions,answers=answers)
        print(default_answer)
        payload, messages = self.set_payload(default_answer, self.gpt_model_name, self.messages)
        output, messages = self.send_response_receive_output(self.url, self.headers, payload, messages)
        self.messages = messages
        self.output=output


    def go_to_step6_2(self):

        self.user_question=self.user_text.get()
        self.question_number_select()

        self.classifier_answer()

        self.GPTanswer=self.output
        self.step6()

    def classifier_answer(self):
        section_num=self.section_num

        questions=[]
        answers=[]
        text=''
        if section_num == 1:
            questions = classifierquestion.render(section=1, modelname=self.modelname)
            answers = classifiersummary1.render(r2=self.accuracy, modelName=self.modelname)
        elif section_num == 2:
            text = ''
            for target_names, row in self.coeff_df.iterrows():
                question = classifierquestion.render(section=2, classes=target_names, ycol=self.selected_dependent_var)
                positivecol = []
                negativecol = []
                notchangecol = []
                i=0
                for self.selected_independent_vars, coeff in row.items():
                    if i!=0:
                        if float(coeff) > 0:
                            positivecol.append(self.selected_independent_vars)
                        elif float(coeff) < 0:
                            negativecol.append(self.selected_independent_vars)
                        else:
                            notchangecol.append(self.selected_independent_vars)
                    i=i+1

                answer = classifiersummary2.render(classes=target_names, pc=positivecol, nc=negativecol,
                                                   ncc=notchangecol)
                questions.append(question)
                text=text + answer + '\n'
            answers = text
        elif section_num == 3:
            feature_importance = np.mean(np.abs(self.coefficients), axis=0)
            questions = classifierquestion.render(section=3, ycol=self.selected_dependent_var)
            most_important_x = self.coeff_df.columns[np.argmax(np.abs(feature_importance))]
            answer = classifiersummary3.render(imp=most_important_x)
            if len(self.target_class)> 2:
                for target_name, row in self.coeff_df.iterrows():
                    most_important_feature = row.abs().idxmax()
                    text = text + f"\nFor the {target_name},the most important feature is {most_important_feature}."
                answers = text + '\n' + answer
            else:
                answers=answer

        elif section_num == 4:
            questions = MLclassifierQuestionSet.render(section=4)
            answers = MachineLearningclassifierSummary1.render(trainR2=self.train_accuracy, testR2=self.test_accuracy)

        elif section_num == 5:
            questions = MLclassifierQuestionSet.render(section=5)

            max_value = self.coeff_df.abs().max().max()
            most_important_x = self.coeff_df.columns[self.coeff_df.abs().iloc[0] == max_value][0]

            answers = MachineLearningclassifierSummary2.render(Xcol=most_important_x, ycol=self.selected_dependent_var,coeff=max_value)

        elif section_num == 0:
            questions="There are no matching default questions in the template."
            answers="Please answer the question based on the analysis results."

        default_answer = answerup.render(userquestion=self.user_text.get(), questions=questions, answers=answers)

        print(default_answer)
        payload, messages = self.set_payload(default_answer, self.gpt_model_name, self.messages)
        output, messages = self.send_response_receive_output(self.url, self.headers, payload, messages)

        self.messages = messages
        self.output=output


    def question_number_select(self):
        # if self.choice.get() == 1 or self.choice.get() ==3:
        #     loader = TextLoader('./apptemplates/linearQuestionset.txt')
        #
        # elif self.choice.get() == 2 or self.choice.get() ==4:
        #     loader = TextLoader('./apptemplates/classifierQuestionset.txt')
        # index = VectorstoreIndexCreator().from_loaders([loader])
        # query ="My question is: "+self.user_question+"\n Please refer to the text and choose the Section number that matches the meaning of my question. Please note that as long as the meaning matches, there is no need for word-for-word correspondence. My entry may have spelling or grammatical mistakes, please ignore those mistakes. Returns 0 if no section matches. Only answer an integer as your choose, do not reply with any information other than the integer, do not reply why you choose the section number."
        # output = index.query(query)

        if self.choice.get() == 1:
            if self.model_var.get()<5:
                with open('./apptemplates/QuestionBank/linearQuestionBank.txt', 'r') as file:
                    content = file.read()
            else:
                with open('./apptemplates/QuestionBank/MLlinearQuestionBank.txt', 'r') as file:
                    content = file.read()
        elif self.choice.get() == 3:
            with open('./apptemplates/QuestionBank/linearQuestionBank.txt', 'r') as file:
                content = file.read()

        elif self.choice.get() == 2:
            if self.model_var.get()<4:
                with open('./apptemplates/QuestionBank/classifierQuestionBank.txt', 'r') as file:
                    content = file.read()
            else:
                with open('./apptemplates/QuestionBank/MLclassifierQuestionBank.txt', 'r') as file:
                    content = file.read()
        elif self.choice.get() == 4:
            with open('./apptemplates/QuestionBank/classifierQuestionBank.txt', 'r') as file:
                content = file.read()

        query = "My question is: " + self.user_question + "\nPlease refer to the following question bank and choose the Section number that matches the meaning of my question. Please note that as long as the meaning matches, there is no need for word-for-word correspondence. My entry may have spelling or grammatical mistakes, please ignore those mistakes. Returns 0 if no section matches. Only answer an integer as you choose, do not reply with any information other than the integer, do not reply why you chose the section number. Following is the question bank: \n"+content

        print(query)
        payload, messages = self.set_payload(query, self.gpt_model_name, self.messages)
        output, messages = self.send_response_receive_output(self.url, self.headers, payload, messages)

        self.section_num = extract_first_integer(output)
        print(output)
        print(self.section_num)

    def set_chatGPT(self, Xcol, ycol, modelname, modelinformation, key, chatmodel="gpt-3.5-turbo", url="https://api.openai.com/v1/chat/completions"):
        openai.api_key = key
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {openai.api_key}"
        }
        background = databackground.render(Xcol=Xcol, ycol=ycol, modelname=modelname, modelinformation=modelinformation)
        messages = [{"role": "system", "content": background}, ]
        return (url, background, chatmodel, headers, messages)

    def set_payload(self, message,GPTmodelname="gpt-3.5-turbo", messages=[]):
        messages.append({"role": "user", "content": message}, )
        payload = {
            "model": GPTmodelname,
            "messages": messages,
            "temperature": 1.0,
            "top_p": 1.0,
            "n": 1,
            "stream": False,
            "presence_penalty": 0,
            "frequency_penalty": 0,
        }
        return (payload, messages)

    def send_response_receive_output(self, URL, headers, payload, messages):
        response = requests.post(URL, headers=headers, json=payload, stream=False)
        print(json.loads(response.content))
        output = json.loads(response.content)["choices"][0]['message']['content']
        messages.append({"role": "assistant", "content": output})
        return (output, messages)

    def save_chat_history_to_docx(self):
        doc = Document()

        for message in self.messages:
            role = message['role']
            content = message['content']

            if role == 'user':
                doc.add_paragraph('User: ' + content)
            elif role == 'assistant':
                doc.add_paragraph('Assistant: ' + content)
            else:
                doc.add_paragraph(content)
        filename="data_report.docx"
        doc.save(filename)

    def get_selected_independent_vars(self):
        return [self.independent_vars.get(i) for i in self.independent_vars.curselection()]

    def cleanData(self, data, threshold, Xcol=[], ycol=''):
        """This function takes in as input a dataset, and returns a clean dataset.

        :param data: This is the dataset that will be cleaned.
        :param treshold: This is the treshold that decides whether columns are deleted or their missing values filled.
        :return: A dataset that does not have any missing values.
        """
        if Xcol != [] and ycol != '':
            data = data[Xcol + [ycol]]
        data = data.replace('?', np.nan)
        data = data.loc[:, data.isnull().mean() < threshold]  # filter data
        imputer = SimpleImputer(missing_values=np.nan, strategy='mean')
        for i in data.columns:
            imputer = imputer.fit(data[[i]])
            data[[i]] = imputer.transform(data[[i]])
        return data
    def more_readable_model_name(self):
        modeldetail=self.modeldetail
        if "Ridge" in modeldetail and "BayesianRidge" not in modeldetail:
            translatedmodel = "Ridge Model"
        elif "LinearDiscriminant" in modeldetail:
            translatedmodel = "Linear Discriminant Analysis"
        elif "GradientBoosting" in modeldetail:
            translatedmodel = "Gradient Boosting Model"
        elif "AdaBoost" in modeldetail:
            translatedmodel = "Ada Boost"
        elif "LGBMClassifier" in modeldetail:
            translatedmodel = "Light Gradient Boosting Machine Classifier"
        elif "DummyClassifier" in modeldetail:
            translatedmodel = "Dummy Classifier"
        elif "KNeighborsClassifier" in modeldetail:
            translatedmodel = "K Neighbors Classifier"
        elif "SGDClassifier" in modeldetail:
            translatedmodel = "SGD Classifier"
        elif "LGBMRegressor" in modeldetail:
            translatedmodel = "Light Gradient Boosting Machine"
        elif "RandomForest" in modeldetail:
            translatedmodel = "Random Forest Model"
        elif "XGBRegressor" in modeldetail:
            translatedmodel = "Extreme Gradient Boosting"
        elif "XGBClassifier" in modeldetail:
            translatedmodel = "Extreme Gradient Boosting Classifier"
        elif "Logistic" in modeldetail:
            translatedmodel = "Logistic Model"
        elif "QuadraticDiscriminant" in modeldetail:
            translatedmodel = "Quadratic Discriminant Analysis"
        elif "GaussianNB" in modeldetail:
            translatedmodel = "Naive Bayes"
        elif "ExtraTrees" in modeldetail:
            translatedmodel = "Extra Trees model"
        elif "DecisionTree" in modeldetail:
            translatedmodel = "Decision Tree Model"
        elif "Lasso" in modeldetail and "LassoLars" not in modeldetail:
            translatedmodel = "Lasso Regression	"
        elif "LassoLars" in modeldetail:
            translatedmodel = "Lasso Least Angle Regression	"
        elif "BayesianRidge" in modeldetail:
            translatedmodel = "Bayesian Ridge"
        elif "LinearRegression" in modeldetail:
            translatedmodel = "Linear Regression"
        elif "HuberRegressor" in modeldetail:
            translatedmodel = "Huber Regressor"
        elif "PassiveAggressiveRegressor" in modeldetail:
            translatedmodel = "Passive Aggressive Regressor"
        elif "OrthogonalMatchingPursuit" in modeldetail:
            translatedmodel = "Orthogonal Matching Pursuit"
        elif "AdaBoostRegressor" in modeldetail:
            translatedmodel = "AdaBoost Regressor"
        elif "KNeighborsRegressor" in modeldetail:
            translatedmodel = "K Neighbors Regressor"
        elif "ElasticNet" in modeldetail:
            translatedmodel = "Elastic Net"
        elif "DummyRegressor" in modeldetail:
            translatedmodel = "Dummy Regressor"
        elif "Lars" in modeldetail:
            translatedmodel = "Least Angle Regression"
        self.modelname=translatedmodel

if __name__ == '__main__':
    app = MyApp()
    app.start()
